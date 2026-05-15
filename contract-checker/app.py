import os
import io
import json
import uuid
import re
import datetime
from pathlib import Path
from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import anthropic

app = Flask(__name__)
app.secret_key = os.urandom(24)

BASE_DIR = Path(__file__).parent

with open(BASE_DIR / "standard_requirements.json", "r", encoding="utf-8") as f:
    STANDARD_REQUIREMENTS = json.load(f)

sessions: dict = {}

client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY", ""))


# ──────────────────────────────────────────────
# Document extraction
# ──────────────────────────────────────────────

def extract_text(file_bytes: bytes, ext: str) -> str:
    if ext == ".docx":
        doc = Document(io.BytesIO(file_bytes))
        lines = []
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip():
                lines.append(f"[P{i}] {p.text}")
        return "\n".join(lines)
    return file_bytes.decode("utf-8", errors="replace")


# ──────────────────────────────────────────────
# Claude analysis
# ──────────────────────────────────────────────

def build_requirements_summary() -> str:
    items = []
    for cat in STANDARD_REQUIREMENTS["categories"]:
        for req in cat["requirements"]:
            items.append({
                "id": req["id"],
                "category": cat["name"],
                "name": req["name"],
                "description": req["description"],
                "mandatory": req["mandatory"],
                "keywords": req["keywords"],
                "standard_wording": req["standard_wording"],
            })
    return json.dumps(items, ensure_ascii=False, indent=2)


def analyze_contract(contract_text: str) -> dict:
    reqs = build_requirements_summary()

    prompt = f"""Ты опытный юрист-эксперт по договорному праву РФ. Проанализируй договор поставки оборудования/материалов для производственного предприятия.

СТАНДАРТНЫЕ ТРЕБОВАНИЯ (обязательные условия заказчика-покупателя):
{reqs}

ТЕКСТ ДОГОВОРА:
{contract_text[:10000]}

Задача: для каждого стандартного требования определить — присутствует в договоре (compliant), изменено/ухудшено (modified), или отсутствует (missing).

Верни ТОЛЬКО валидный JSON (без markdown, без пояснений вне JSON):
{{
  "contract_requisites": {{
    "number": "номер договора или 'не указан'",
    "date": "дата договора или 'не указана'",
    "supplier": "полное наименование поставщика",
    "buyer": "полное наименование покупателя",
    "subject": "краткое описание предмета договора"
  }},
  "missing_clauses": [
    {{
      "requirement_id": "REQ-XX-XX",
      "requirement_name": "название требования",
      "description": "что именно отсутствует и чем это опасно для интересов заказчика",
      "recommended_wording": "готовая юридическая формулировка для включения в договор"
    }}
  ],
  "modified_clauses": [
    {{
      "requirement_id": "REQ-XX-XX",
      "requirement_name": "название требования",
      "contract_wording": "точная цитата из договора (не более 200 символов)",
      "issue": "в чем конкретно ухудшение или несоответствие интересам заказчика",
      "recommended_correction": "рекомендуемая исправленная формулировка"
    }}
  ],
  "compliant_clauses": [
    {{
      "requirement_id": "REQ-XX-XX",
      "requirement_name": "название требования",
      "contract_wording": "краткая цитата соответствующего условия из договора"
    }}
  ],
  "risk_level": "высокий|средний|низкий",
  "overall_assessment": "общая оценка договора в 2-3 предложения с позиции интересов заказчика",
  "key_risks": ["конкретный риск 1 для заказчика", "риск 2", "риск 3"]
}}"""

    message = client.messages.create(
        model="claude-opus-4-7",
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}],
    )

    raw = message.content[0].text.strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)

    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        m = re.search(r"\{[\s\S]*\}", raw)
        if m:
            return json.loads(m.group())
        raise ValueError("Не удалось разобрать ответ модели")


# ──────────────────────────────────────────────
# Document generation helpers
# ──────────────────────────────────────────────

def _page_setup(doc: Document) -> None:
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(3)
        s.right_margin = Cm(1.5)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)


def _body(doc: Document, text: str, indent: float = 0, color: tuple = None,
          italic: bool = False, size: int = 11, space_after: int = 4) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def _contract_ref(analysis: dict) -> str:
    req = analysis.get("contract_requisites", {})
    num = req.get("number", "б/н")
    date = req.get("date", "б/д")
    return f"№ {num} от {date}"


# ──────────────────────────────────────────────
# Generate conclusion
# ──────────────────────────────────────────────

def gen_conclusion(analysis: dict) -> bytes:
    doc = Document()
    _page_setup(doc)
    req = analysis.get("contract_requisites", {})
    today = datetime.date.today().strftime("%d.%m.%Y")
    ref = _contract_ref(analysis)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(2)
    r = t.add_run("ЮРИДИЧЕСКОЕ ЗАКЛЮЧЕНИЕ")
    r.bold = True; r.font.size = Pt(14)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.paragraph_format.space_after = Pt(14)
    r2 = t2.add_run(f"по результатам правовой экспертизы\nДоговора поставки {ref}")
    r2.font.size = Pt(11); r2.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    _body(doc, f"Дата подготовки: {today}", size=10)

    # 1. Реквизиты
    _heading(doc, "1. РЕКВИЗИТЫ ДОГОВОРА")
    for lbl, val in [
        ("Поставщик:", req.get("supplier", "не определен")),
        ("Покупатель (Заказчик):", req.get("buyer", "не определен")),
        ("Предмет:", req.get("subject", "не определен")),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        rb = p.add_run(f"{lbl} "); rb.bold = True; rb.font.size = Pt(11)
        p.add_run(val).font.size = Pt(11)

    risk = analysis.get("risk_level", "не определен")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    rb = p.add_run("Уровень риска для Заказчика: "); rb.bold = True; rb.font.size = Pt(11)
    rv = p.add_run(risk.upper())
    rv.font.size = Pt(11); rv.bold = True
    color_map = {"высокий": (192, 57, 43), "средний": (214, 122, 0), "низкий": (39, 130, 67)}
    rv.font.color.rgb = RGBColor(*color_map.get(risk.lower(), (50, 50, 50)))

    # 2. Общая оценка
    _heading(doc, "2. ОБЩАЯ ОЦЕНКА")
    _body(doc, analysis.get("overall_assessment", ""))

    # 3. Итоги экспертизы
    missing = analysis.get("missing_clauses", [])
    modified = analysis.get("modified_clauses", [])
    compliant = analysis.get("compliant_clauses", [])
    total_issues = len(missing) + len(modified)

    _heading(doc, "3. РЕЗУЛЬТАТЫ ЭКСПЕРТИЗЫ")
    for txt, cnt, col in [
        (f"Отсутствующих обязательных условий: {len(missing)}", len(missing), (192, 57, 43)),
        (f"Условий с замечаниями (изменённые формулировки): {len(modified)}", len(modified), (214, 122, 0)),
        (f"Условий, соответствующих требованиям Заказчика: {len(compliant)}", len(compliant), (39, 130, 67)),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"• {txt}")
        run.font.size = Pt(11)
        if cnt > 0:
            run.font.color.rgb = RGBColor(*col)

    # 4. Отсутствующие условия
    sec = 4
    if missing:
        _heading(doc, f"{sec}. ОТСУТСТВУЮЩИЕ ОБЯЗАТЕЛЬНЫЕ УСЛОВИЯ")
        for i, cl in enumerate(missing, 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
            r = p.add_run(f"{i}. {cl['requirement_name']} ({cl['requirement_id']})")
            r.bold = True; r.font.size = Pt(11)
            _body(doc, f"Замечание: {cl['description']}", indent=0.25, size=10, space_after=2)
            _body(doc, f"Рекомендуемая формулировка: {cl['recommended_wording']}",
                  indent=0.25, size=10, italic=True, color=(39, 130, 67), space_after=6)
        sec += 1

    # 5. Изменённые условия
    if modified:
        _heading(doc, f"{sec}. УСЛОВИЯ С ЗАМЕЧАНИЯМИ")
        for i, cl in enumerate(modified, 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
            r = p.add_run(f"{i}. {cl['requirement_name']} ({cl['requirement_id']})")
            r.bold = True; r.font.size = Pt(11)
            if cl.get("contract_wording"):
                _body(doc, f"В договоре: «{cl['contract_wording']}»", indent=0.25, size=10,
                      italic=True, space_after=2)
            _body(doc, f"Замечание: {cl['issue']}", indent=0.25, size=10,
                  color=(192, 57, 43), space_after=2)
            _body(doc, f"Рекомендация: {cl['recommended_correction']}", indent=0.25, size=10,
                  italic=True, color=(39, 130, 67), space_after=6)
        sec += 1

    # Key risks
    risks = analysis.get("key_risks", [])
    if risks:
        _heading(doc, f"{sec}. КЛЮЧЕВЫЕ РИСКИ ДЛЯ ЗАКАЗЧИКА")
        for risk_text in risks:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(f"• {risk_text}")
            r.font.size = Pt(11); r.font.color.rgb = RGBColor(192, 57, 43)

    # Вывод
    doc.add_paragraph()
    _heading(doc, "ВЫВОД")
    if total_issues == 0:
        _body(doc, f"Договор поставки {ref} в целом соответствует требованиям Заказчика. "
              "Замечаний по обязательным условиям не выявлено.")
    else:
        _body(doc, f"Договор поставки {ref} требует доработки: выявлено {total_issues} "
              f"замечание(-ий). Рекомендуется подготовить протокол разногласий или "
              "дополнительное соглашение в соответствии с настоящим Заключением.")

    # Подпись
    doc.add_paragraph()
    _body(doc, "Юрист: ________________________  /________________________/")
    _body(doc, f"Дата: {today}", size=10)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


# ──────────────────────────────────────────────
# Generate protocol of disagreements
# ──────────────────────────────────────────────

def gen_protocol(analysis: dict) -> bytes:
    doc = Document()
    _page_setup(doc)
    req = analysis.get("contract_requisites", {})
    today = datetime.date.today().strftime("%d.%m.%Y")
    ref = _contract_ref(analysis)
    supplier = req.get("supplier", "[Поставщик]")
    buyer = req.get("buyer", "[Покупатель]")

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(2)
    r = t.add_run("ПРОТОКОЛ РАЗНОГЛАСИЙ")
    r.bold = True; r.font.size = Pt(14)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.paragraph_format.space_after = Pt(12)
    r2 = t2.add_run(f"к Договору поставки {ref}")
    r2.bold = True; r2.font.size = Pt(12)

    dp = doc.add_paragraph()
    dp.paragraph_format.space_after = Pt(12)
    dp.add_run(
        f"г. _________________________                                          "
        f"«___» _____________ {datetime.date.today().year} г."
    ).font.size = Pt(11)

    pre = doc.add_paragraph()
    pre.paragraph_format.space_after = Pt(12)
    pre.add_run(
        f"{buyer}, именуемый в дальнейшем «Покупатель», в лице "
        f"________________________________, действующего на основании ________________, "
        f"с одной стороны, и {supplier}, именуемый в дальнейшем «Поставщик», в лице "
        f"________________________________, действующего на основании ________________, "
        f"с другой стороны, составили настоящий Протокол разногласий к Договору поставки "
        f"{ref} о нижеследующем:"
    ).font.size = Pt(11)

    missing = analysis.get("missing_clauses", [])
    modified = analysis.get("modified_clauses", [])

    # Table
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdrs = ["№\nп/п", f"Редакция\nПоставщика\n(по договору)",
            f"Редакция\nПокупателя\n(предлагается)", "Основание / пояснение"]
    for cell, h in zip(tbl.rows[0].cells, hdrs):
        cell.text = h
        r = cell.paragraphs[0].runs[0]
        r.bold = True; r.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    row_num = 1
    for cl in modified:
        row = tbl.add_row()
        cells = row.cells
        cells[0].text = str(row_num)
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[1].text = cl.get("contract_wording", "[формулировка договора]")
        cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        corr = cl.get("recommended_correction", "")
        cells[2].text = corr
        cells[2].paragraphs[0].runs[0].font.size = Pt(9)
        cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 100, 0)
        cells[3].text = f"{cl['requirement_name']}: {cl.get('issue', '')}"
        cells[3].paragraphs[0].runs[0].font.size = Pt(9)
        row_num += 1

    for cl in missing:
        row = tbl.add_row()
        cells = row.cells
        cells[0].text = str(row_num)
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[1].text = "Отсутствует в Договоре"
        cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(180, 40, 30)
        wording = cl.get("recommended_wording", "")
        cells[2].text = wording
        cells[2].paragraphs[0].runs[0].font.size = Pt(9)
        cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 100, 0)
        cells[3].text = f"Дополнить: {cl['requirement_name']}. {cl.get('description', '')}"
        cells[3].paragraphs[0].runs[0].font.size = Pt(9)
        row_num += 1

    # Set column widths
    widths = [Cm(1.2), Cm(5.5), Cm(6.5), Cm(4)]
    for row in tbl.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w

    doc.add_paragraph()
    _body(doc,
          f"Настоящий Протокол разногласий является неотъемлемой частью Договора поставки {ref}. "
          "Условия Договора, не затронутые настоящим Протоколом, остаются в силе.")

    doc.add_paragraph()
    st = doc.add_paragraph()
    r = st.add_run("ПОДПИСИ СТОРОН")
    r.bold = True; r.font.size = Pt(11)

    sig_tbl = doc.add_table(rows=5, cols=2)
    sig_data = [
        [f"ПОКУПАТЕЛЬ\n{buyer}", f"ПОСТАВЩИК\n{supplier}"],
        ["", ""],
        ["________________________________", "________________________________"],
        ["(подпись / расшифровка)", "(подпись / расшифровка)"],
        [f"М.П.        Дата: {today}", f"М.П.        Дата: {today}"],
    ]
    for i, row_d in enumerate(sig_data):
        for j, txt in enumerate(row_d):
            cell = sig_tbl.rows[i].cells[j]
            cell.text = txt
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(10)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


# ──────────────────────────────────────────────
# Generate additional agreement
# ──────────────────────────────────────────────

def gen_agreement(analysis: dict) -> bytes:
    doc = Document()
    _page_setup(doc)
    req = analysis.get("contract_requisites", {})
    today = datetime.date.today().strftime("%d.%m.%Y")
    ref = _contract_ref(analysis)
    supplier = req.get("supplier", "[Поставщик]")
    buyer = req.get("buyer", "[Покупатель]")

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(2)
    r = t.add_run("ДОПОЛНИТЕЛЬНОЕ СОГЛАШЕНИЕ")
    r.bold = True; r.font.size = Pt(14)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.paragraph_format.space_after = Pt(12)
    r2 = t2.add_run(f"к Договору поставки {ref}")
    r2.bold = True; r2.font.size = Pt(12)

    dp = doc.add_paragraph()
    dp.paragraph_format.space_after = Pt(12)
    dp.add_run(
        f"г. _________________________                                          "
        f"«___» _____________ {datetime.date.today().year} г."
    ).font.size = Pt(11)

    pre = doc.add_paragraph()
    pre.paragraph_format.space_after = Pt(12)
    pre.add_run(
        f"{buyer}, именуемый в дальнейшем «Покупатель», в лице "
        f"________________________________, действующего на основании ________________, "
        f"с одной стороны, и {supplier}, именуемый в дальнейшем «Поставщик», в лице "
        f"________________________________, действующего на основании ________________, "
        f"с другой стороны, заключили настоящее Дополнительное соглашение к Договору "
        f"поставки {ref} о нижеследующем:"
    ).font.size = Pt(11)

    missing = analysis.get("missing_clauses", [])
    for i, cl in enumerate(missing, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(
            f"{i}. Стороны договорились дополнить Договор поставки {ref} "
            f"следующим условием ({cl['requirement_name']}, {cl['requirement_id']}):"
        )
        r.bold = True; r.font.size = Pt(11)
        _body(doc, f"«{cl['recommended_wording']}»", indent=0.3, italic=True, size=11, space_after=8)

    doc.add_paragraph()
    _body(doc,
          f"Во всем остальном, что не предусмотрено настоящим Дополнительным соглашением, "
          f"Стороны руководствуются условиями Договора поставки {ref}. "
          "Настоящее Дополнительное соглашение составлено в двух экземплярах, имеющих равную "
          "юридическую силу, по одному для каждой из Сторон.")

    doc.add_paragraph()
    st = doc.add_paragraph()
    r = st.add_run("РЕКВИЗИТЫ И ПОДПИСИ СТОРОН")
    r.bold = True; r.font.size = Pt(11)

    sig_tbl = doc.add_table(rows=5, cols=2)
    sig_data = [
        [f"ПОКУПАТЕЛЬ\n{buyer}", f"ПОСТАВЩИК\n{supplier}"],
        ["", ""],
        ["________________________________", "________________________________"],
        ["(подпись / расшифровка)", "(подпись / расшифровка)"],
        [f"М.П.        Дата: {today}", f"М.П.        Дата: {today}"],
    ]
    for i, row_d in enumerate(sig_data):
        for j, txt in enumerate(row_d):
            cell = sig_tbl.rows[i].cells[j]
            cell.text = txt
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(10)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


# ──────────────────────────────────────────────
# Generate checked document (original + margin annotations)
# ──────────────────────────────────────────────

def gen_checked_doc(file_bytes: bytes, analysis: dict) -> bytes:
    doc = Document(io.BytesIO(file_bytes))
    req = analysis.get("contract_requisites", {})
    ref = _contract_ref(analysis)
    today = datetime.date.today().strftime("%d.%m.%Y")

    modified = analysis.get("modified_clauses", [])
    missing = analysis.get("missing_clauses", [])

    # Highlight modified paragraphs in orange
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        para_lower = para.text.lower()
        for cl in modified:
            kws = [w for w in cl.get("requirement_name", "").lower().split() if len(w) > 3]
            if any(kw in para_lower for kw in kws):
                for run in para.runs:
                    run.font.color.rgb = RGBColor(180, 90, 0)
                break

    # Append summary at end
    doc.add_paragraph()
    sep = doc.add_paragraph("─" * 60)
    sep.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    r = p.add_run(f"РЕЗУЛЬТАТЫ ПРАВОВОЙ ЭКСПЕРТИЗЫ — Договор {ref}")
    r.bold = True; r.font.size = Pt(12)

    _body_in_doc(doc, f"Дата проверки: {today}", size=10)
    _body_in_doc(doc, f"Отсутствующих условий: {len(missing)}", size=10,
                 color=(180, 40, 30) if missing else None)
    _body_in_doc(doc, f"Условий с замечаниями: {len(modified)}", size=10,
                 color=(180, 90, 0) if modified else None)

    if missing:
        p = doc.add_paragraph()
        r = p.add_run("ОТСУТСТВУЮЩИЕ УСЛОВИЯ:")
        r.bold = True; r.font.size = Pt(10)
        for cl in missing:
            _body_in_doc(doc, f"• {cl['requirement_name']} ({cl['requirement_id']}): {cl['description']}",
                         size=9, color=(180, 40, 30))

    if modified:
        p = doc.add_paragraph()
        r = p.add_run("УСЛОВИЯ С ЗАМЕЧАНИЯМИ (выделены оранжевым):")
        r.bold = True; r.font.size = Pt(10)
        for cl in modified:
            _body_in_doc(doc, f"• {cl['requirement_name']}: {cl['issue']}", size=9,
                         color=(180, 90, 0))

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def _body_in_doc(doc, text, size=11, color=None, indent=0, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


# ──────────────────────────────────────────────
# Routes
# ──────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/analyze", methods=["POST"])
def api_analyze():
    if "contract" not in request.files:
        return jsonify({"error": "Файл не загружен"}), 400

    file = request.files["contract"]
    if not file.filename:
        return jsonify({"error": "Файл не выбран"}), 400

    ext = Path(file.filename).suffix.lower()
    if ext not in {".docx", ".txt"}:
        return jsonify({"error": "Поддерживаются форматы: .docx, .txt"}), 400

    try:
        file_bytes = file.read()
        contract_text = extract_text(file_bytes, ext)
        if not contract_text.strip():
            return jsonify({"error": "Файл пустой или не содержит текста"}), 400

        result = analyze_contract(contract_text)

        sid = str(uuid.uuid4())
        sessions[sid] = {
            "analysis": result,
            "file_bytes": file_bytes,
            "filename": file.filename,
            "ext": ext,
        }
        return jsonify({"session_id": sid, "analysis": result})

    except anthropic.AuthenticationError:
        return jsonify({"error": "Ошибка аутентификации API. Проверьте ANTHROPIC_API_KEY."}), 500
    except Exception as e:
        return jsonify({"error": f"Ошибка анализа: {str(e)}"}), 500


def _get_session(sid):
    if sid not in sessions:
        return None, ("Сессия не найдена или истекла", 404)
    return sessions[sid], None


def _safe_filename(name: str) -> str:
    return re.sub(r'[\\/*?:"<>|]', "-", name)


@app.route("/api/download/checked/<sid>")
def download_checked(sid):
    data, err = _get_session(sid)
    if err:
        return jsonify({"error": err[0]}), err[1]
    if data["ext"] != ".docx":
        return jsonify({"error": "Доступно только для .docx"}), 400
    try:
        b = gen_checked_doc(data["file_bytes"], data["analysis"])
        name = _safe_filename(f"Проверенный_{data['filename']}")
        return send_file(io.BytesIO(b),
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                         as_attachment=True, download_name=name)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/download/conclusion/<sid>")
def download_conclusion(sid):
    data, err = _get_session(sid)
    if err:
        return jsonify({"error": err[0]}), err[1]
    try:
        b = gen_conclusion(data["analysis"])
        req = data["analysis"].get("contract_requisites", {})
        num = _safe_filename(req.get("number", "бн"))
        name = f"Заключение_к_договору_{num}.docx"
        return send_file(io.BytesIO(b),
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                         as_attachment=True, download_name=name)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/download/protocol/<sid>")
def download_protocol(sid):
    data, err = _get_session(sid)
    if err:
        return jsonify({"error": err[0]}), err[1]
    try:
        b = gen_protocol(data["analysis"])
        req = data["analysis"].get("contract_requisites", {})
        num = _safe_filename(req.get("number", "бн"))
        name = f"Протокол_разногласий_{num}.docx"
        return send_file(io.BytesIO(b),
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                         as_attachment=True, download_name=name)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/download/agreement/<sid>")
def download_agreement(sid):
    data, err = _get_session(sid)
    if err:
        return jsonify({"error": err[0]}), err[1]
    try:
        b = gen_agreement(data["analysis"])
        req = data["analysis"].get("contract_requisites", {})
        num = _safe_filename(req.get("number", "бн"))
        name = f"Доп_соглашение_{num}.docx"
        return send_file(io.BytesIO(b),
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                         as_attachment=True, download_name=name)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    app.run(debug=True, port=5000)
